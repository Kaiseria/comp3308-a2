from docx import Document

import contextlib
import io
import re

import test_dt
import test_dtplus

PERFORMANCE_TABLE_LATEX = ""


def latex_escape(text):
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }

    text = text.replace("’", "'")
    text = text.replace("“", "``")
    text = text.replace("”", "''")
    text = text.replace("–", "-")
    text = text.replace("—", "-")

    escaped = ""
    for char in text:
        escaped += replacements.get(char, char)

    return escaped


def verbatim_block(text):
    return (
        r"{\scriptsize" + "\n"
        + r"\begin{verbatim}" + "\n"
        + text
        + r"\end{verbatim}" + "\n"
        + r"}"
    )


def write_tree_file(filename, title, tree, diagram_function, leaves_function,
                    size_function):
    with open(filename, "w") as f:
        f.write("=== Classifier model (full training set) ===\n")
        f.write(title + "\n")
        f.write("------------------\n\n")
        f.write(diagram_function(tree))
        f.write("\n")
        f.write("Number of Leaves  :\t" + str(leaves_function(tree)) + "\n")
        f.write("Size of the tree :\t" + str(size_function(tree)) + "\n")


def parse_weka_results(filename):
    document = Document(filename)
    paragraphs = [p.text.strip() for p in document.paragraphs]
    blocks = []
    current = []

    for text in paragraphs:
        if text == "=== Run information ===":
            if current != []:
                blocks.append(current)
            current = [text]
        elif current != []:
            current.append(text)

    if current != []:
        blocks.append(current)

    results = {}

    for block in blocks:
        scheme = next((x for x in block if x.startswith("Scheme:")), "")
        name = weka_name_from_scheme(scheme)

        if name is None:
            continue

        correct_line = next(
            (x for x in block if x.startswith("Correctly Classified Instances")),
            ""
        )
        weighted_line = next(
            (x for x in block if x.startswith("Weighted Avg.")),
            ""
        )

        accuracy_match = re.search(r"(\d+\.\d+)\s*%", correct_line)
        accuracy = accuracy_match.group(1) if accuracy_match else "-"

        parts = weighted_line.split()
        precision = parts[4] if len(parts) > 4 else "-"
        recall = parts[5] if len(parts) > 5 else "-"
        f1 = parts[6] if len(parts) > 6 else "-"

        if precision == "?":
            precision = "-"
        if recall == "?":
            recall = "-"
        if f1 == "?":
            f1 = "-"

        results[name] = {
            "accuracy": accuracy,
            "precision": precision,
            "recall": recall,
            "f1": f1,
        }

    return results


def weka_name_from_scheme(scheme):
    if "ZeroR" in scheme:
        return "ZeroR"
    if "OneR" in scheme:
        return "1R"
    if "IBk" in scheme:
        return "KNN"
    if "NaiveBayes" in scheme:
        return "NB"
    if "MultilayerPerceptron" in scheme:
        return "MLP"
    if "SMO" in scheme:
        return "SVM"
    if "J48 -U" in scheme:
        return "DT unpruned"
    if "J48 -C" in scheme:
        return "DT pruned"
    if "Bagging" in scheme:
        return "Bagg"
    if "AdaBoostM1" in scheme:
        return "Boost"
    if "RandomForest" in scheme:
        return "RF"

    return None


def read_folds_for_knn(folds_filename):
    folds = []
    current_fold = []

    with open(folds_filename, "r", errors="replace") as f:
        reader = csv_reader(f)

        for row in reader:
            row = [x.strip().replace("\x00", "") for x in row]

            if len(row) == 0 or row[0] == "":
                continue

            if row[0].startswith("fold"):
                if current_fold != []:
                    folds.append(current_fold)
                current_fold = []
            else:
                current_fold.append(row)

        if current_fold != []:
            folds.append(current_fold)

    return folds


def csv_reader(file_object):
    import csv
    return csv.reader(file_object)


def predict_knn(training_rows, testing_row, k, weighted):
    distances = []

    for index, data in enumerate(training_rows):
        distance = 0

        for i in range(len(testing_row)):
            if testing_row[i] != data[i]:
                distance += 1

        distances.append([distance, index, data[-1]])

    distances.sort(key=lambda x: (x[0], x[1]))
    nearest = distances[:k]

    if weighted:
        died_score = 0
        survived_score = 0

        for item in nearest:
            weight = 1 / (item[0] + 0.25)

            if item[-1] == "died":
                died_score += weight
            else:
                survived_score += weight

        if survived_score > died_score:
            return "survived"
        return "died"

    died = 0
    survived = 0

    for item in nearest:
        if item[-1] == "died":
            died += 1
        else:
            survived += 1

    if survived > died:
        return "survived"
    return "died"


def evaluate_knn_metrics(folds_filename, k, weighted):
    folds = read_folds_for_knn(folds_filename)
    predictions = []
    actual_labels = []

    for i in range(10):
        training = []

        for j in range(10):
            if j != i:
                training.extend(folds[j])

        for row in folds[i]:
            predictions.append(predict_knn(training, row[:-1], k, weighted))
            actual_labels.append(row[-1])

    return test_dt.evaluate_predictions(predictions, actual_labels)


def metric_dict_from_evaluation(metrics):
    return {
        "accuracy": f"{metrics['accuracy'] * 100:.4f}",
        "precision": f"{metrics['weighted_precision']:.3f}",
        "recall": f"{metrics['weighted_recall']:.3f}",
        "f1": f"{metrics['weighted_f1']:.3f}",
    }


def build_performance_results(folds_filename):
    results = parse_weka_results("Run information.docx")

    with contextlib.redirect_stdout(io.StringIO()):
        mydt_metrics = test_dt.evaluate_dt_10fold(folds_filename)
        mydtplus_metrics = test_dtplus.evaluate_dt_plus_10fold(folds_filename)

    results["MyKNN"] = metric_dict_from_evaluation(
        evaluate_knn_metrics(folds_filename, 3, False)
    )
    results["MyKNN+"] = metric_dict_from_evaluation(
        evaluate_knn_metrics(folds_filename, 6, True)
    )
    results["MyDT"] = metric_dict_from_evaluation(mydt_metrics)
    results["MyDT+"] = metric_dict_from_evaluation(mydtplus_metrics)

    return results


def make_performance_table(results):
    first = ["ZeroR", "1R", "KNN", "NB", "MLP", "SVM", "MyKNN", "MyKNN+"]
    second = [
        "DT unpruned",
        "DT pruned",
        "MyDT",
        "MyDT+",
        "Bagg",
        "Boost",
        "RF",
    ]

    return (
        r"\textbf{Predictive performance -- accuracy and other suitable measures}"
        + "\n\n"
        + make_metric_table(first, results)
        + "\n\n"
        + make_metric_table(second, results)
        + "\n\n"
        + "The values for the Weka classifiers are taken from "
        + r"\texttt{Run information.docx}. The values for MyKNN, MyKNN+, MyDT and MyDT+ "
        + "were regenerated from the submitted Python classifiers."
    )


def make_metric_table(classifiers, results):
    alignment = "|l|" + "c|" * len(classifiers)
    output = []

    output.append(r"{\scriptsize")
    output.append(r"\begin{center}")
    output.append(r"\begin{tabular}{" + alignment + r"}")
    output.append(r"\hline")
    output.append("Measure & " + " & ".join(classifiers) + r" \\")
    output.append(r"\hline")

    rows = [
        ("Accuracy (\\%)", "accuracy"),
        ("Weighted Precision", "precision"),
        ("Weighted Recall", "recall"),
        ("Weighted F1", "f1"),
    ]

    for label, key in rows:
        values = []

        for classifier in classifiers:
            values.append(results.get(classifier, {}).get(key, "-"))

        output.append(label + " & " + " & ".join(values) + r" \\")
        output.append(r"\hline")

    output.append(r"\end{tabular}")
    output.append(r"\end{center}")
    output.append(r"}")

    return "\n".join(output)


def section_for_index(index, text):
    if index == 0:
        return r"\section{Introduction}"
    if index == 2:
        return r"\section{Data}"
    if text == "KNN+":
        return r"\section{The KNN+ and DT+ Algorithms}" + "\n" + r"\subsection{KNN+}"
    if text == "DT+":
        return r"\subsection{DT+}"
    if index == 26:
        return (
            r"\section{Results and Discussion}"
            + "\n"
            + r"\subsection{Results}"
            + "\n\n"
            + PERFORMANCE_TABLE_LATEX
            + "\n\n"
            + "\n"
            + r"\subsubsection{Accuracy}"
        )
    if index == 27:
        return r"\subsubsection{Precision}"
    if index == 31:
        return r"\subsubsection{Recall}"
    if index == 34:
        return r"\subsection{Discussion}" + "\n" + r"\subsubsection{Comparison of MyKNN, MyKNN+, MyDT and MyDT+}"
    if index == 44:
        return r"\subsubsection{Comparison with Weka Classifiers}"
    if index == 48:
        return r"\subsubsection{Effect of Pruning in Weka Decision Trees}"
    if index == 54:
        return r"\subsubsection{Comparison of Weka Tree-Based Classifiers}"
    if text == "5. Conclusion":
        return r"\section{Conclusion}"

    return None


def paragraph_to_latex(index, text):
    heading = section_for_index(index, text)

    if heading is not None:
        if text in ["KNN+", "DT+", "5. Conclusion"]:
            return heading
        return heading + "\n\n" + latex_escape(text)

    if text == "weight = 1 / (distance + 0.25)":
        return r"\[" + "\n" + r"\mathrm{weight} = \frac{1}{\mathrm{distance} + 0.25}" + "\n" + r"\]"

    return latex_escape(text)


def read_file(filename):
    with open(filename, "r") as f:
        return f.read()


def main():
    global PERFORMANCE_TABLE_LATEX

    folds_file = "heart-folds.csv"
    PERFORMANCE_TABLE_LATEX = make_performance_table(
        build_performance_results(folds_file)
    )
    training_rows = test_dt.all_training_rows_from_folds(folds_file)

    mydt_tree = test_dt.train_dt(training_rows)
    mydtplus_tree = test_dtplus.train_dt_plus(training_rows)

    write_tree_file(
        "mydt_tree.txt",
        "MyDT tree",
        mydt_tree,
        test_dt.print_tree_diagram,
        test_dt.count_leaves,
        test_dt.tree_size,
    )
    write_tree_file(
        "mydtplus_tree.txt",
        "MyDT+ tree",
        mydtplus_tree,
        test_dtplus.print_tree_diagram,
        test_dtplus.count_leaves,
        test_dtplus.tree_size,
    )

    document = Document("a2report.docx")
    body_parts = []

    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()

        if text == "":
            continue

        body_parts.append(paragraph_to_latex(index, text))

    mydt_tree_text = read_file("mydt_tree.txt")
    mydtplus_tree_text = read_file("mydtplus_tree.txt")
    knnplus_code = read_file("test_knnplus.py")
    dtplus_code = read_file("test_dtplus.py")

    tex = r"""\documentclass[12pt,a4paper]{article}
\setlength{\oddsidemargin}{0mm}
\setlength{\evensidemargin}{0mm}
\setlength{\textwidth}{160mm}
\setlength{\topmargin}{-10mm}
\setlength{\textheight}{240mm}
\setlength{\parindent}{0pt}
\setlength{\parskip}{1\baselineskip}
\linespread{1.2}

\title{COMP3308 Assignment 2 Part 2: Report}
\author{Student 1: \textless SID\textgreater \\ Student 2: \textless SID\textgreater}
\date{}

\begin{document}
\maketitle

""" + "\n\n".join(body_parts) + r"""

\section{Reflection}

\textbf{Student 1: \textless SID\textgreater}

Write your reflection here.

\textbf{Student 2: \textless SID\textgreater}

Write your reflection here.

\section{Acknowledgement of AI Use}

\textbf{Student 1: \textless SID\textgreater}

Write your AI-use acknowledgement here, or state: No AI tools were used in this assignment.

\textbf{Student 2: \textless SID\textgreater}

Write your AI-use acknowledgement here, or state: No AI tools were used in this assignment.

\appendix
\section*{Appendix A: Code for KNN+}
""" + verbatim_block(knnplus_code) + r"""

\section*{Appendix B: Code for DT+}
""" + verbatim_block(dtplus_code) + r"""

\section*{Appendix C: MyDT Decision Tree Diagram}
""" + verbatim_block(mydt_tree_text) + r"""

\section*{Appendix D: MyDT+ Decision Tree Diagram}
""" + verbatim_block(mydtplus_tree_text) + r"""

\end{document}
"""

    with open("report.tex", "w") as f:
        f.write(tex)


if __name__ == "__main__":
    main()
